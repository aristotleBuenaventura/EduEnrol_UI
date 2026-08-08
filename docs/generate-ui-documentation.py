#!/usr/bin/env python3
"""
Generates EduEnrol UI Technical Documentation (Word format).
Run from repo root: python3 docs/generate-ui-documentation.py
"""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SCREENSHOTS = Path(__file__).resolve().parent / "screenshots"
OUTPUT = Path(__file__).resolve().parent / "EduEnrol-UI-Technical-Documentation.docx"

DOC_VERSION = "1.1"
DOC_DATE = date(2026, 5, 17).strftime("%d %B %Y")


def set_document_defaults(doc):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15


def add_page_break(doc):
    doc.add_page_break()


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_bullet(doc, text):
    return doc.add_paragraph(text, style="List Bullet")


def add_numbered(doc, text):
    return doc.add_paragraph(text, style="List Number")


def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row_data in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row_data):
            cells[i].text = str(val)
    doc.add_paragraph()
    return table


def add_image(doc, filename, caption, width_inches=6.2):
    path = SCREENSHOTS / filename
    if not path.exists():
        add_para(doc, f"[Screenshot: {filename}]", italic=True)
        return
    doc.add_picture(str(path), width=Inches(width_inches))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph()


def add_table_of_contents(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    for el_type in ("begin", "separate", "end"):
        fld = OxmlElement("w:fldChar")
        fld.set(qn("w:fldCharType"), el_type if el_type != "separate" else "separate")
        if el_type == "begin":
            run._r.append(fld)
            instr = OxmlElement("w:instrText")
            instr.set(qn("xml:space"), "preserve")
            instr.text = 'TOC \\o "1-3" \\h \\z \\u'
            run._r.append(instr)
        elif el_type == "separate":
            run._r.append(fld)
        else:
            run._r.append(fld)
    add_para(
        doc,
        'After opening in Word: right-click the table of contents → "Update Field" → "Update entire table".',
        italic=True,
    )


def build_document():
    doc = Document()
    set_document_defaults(doc)

    # Cover
    for _ in range(5):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t = title.add_run("EduEnrol User Interface\nTechnical Documentation")
    t.bold = True
    t.font.size = Pt(26)
    t.font.color.rgb = RGBColor(0x16, 0x3A, 0x5C)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = sub.add_run("Front-end architecture, tools, and component reuse")
    s.font.size = Pt(13)
    s.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in [
        f"Version {DOC_VERSION}  |  {DOC_DATE}",
        "Prepared by: UI Development Team",
        "Classification: Client technical reference",
    ]:
        meta.add_run(line + "\n").font.size = Pt(11)
    add_page_break(doc)

    add_heading(doc, "Table of Contents", 1)
    add_table_of_contents(doc)
    add_page_break(doc)

    # 1 Executive summary
    add_heading(doc, "1. Executive Summary", 1)
    add_para(
        doc,
        "EduEnrol is a single-page application (SPA) built with React. The interface is organised "
        "around four role-based portals—Parent, Administrator, SLT (school leadership team), and "
        "Enrolment Manager—plus two public login screens. Rather than building four separate "
        "applications, one shared application shell (sidebar, top bar, scrollable content) is "
        "configured per role while pages compose reusable dashboard and form components.",
    )
    add_para(doc, "This document explains in concrete terms:")
    add_bullet(doc, "Which tools were used during development and what each is responsible for")
    add_bullet(doc, "How the UI loads, routes, and renders from browser to screen")
    add_bullet(doc, "How components are designed for reuse across roles and features")
    add_bullet(doc, "Representative screenshots of finished screens")

    # 2 Introduction
    add_heading(doc, "2. Introduction", 1)
    add_heading(doc, "2.1 Purpose and audience", 2)
    add_para(
        doc,
        "The audience for this document includes product owners, technical reviewers, and "
        "future developers who need to understand how the front end is structured—not only "
        "what screens exist, but how navigation, state, styling, and shared widgets fit together.",
    )
    add_heading(doc, "2.2 Delivery approach", 2)
    add_para(
        doc,
        "Implementation followed agreed functional requirements and visual specifications. Work "
        "proceeded screen-by-screen: layout regions were identified first, then shared components "
        "were extracted when the same pattern appeared on a second screen (e.g. status pills, "
        "metric cards, panel wrappers). Dummy data modules under src/data/ keep pages realistic "
        "until back-end APIs are connected.",
    )

    # 3 Tools
    add_heading(doc, "3. Development Tools and Environment", 1)
    add_para(
        doc,
        "The UI was built with a focused toolchain—no heavy UI framework beyond React itself. "
        "The table below lists each tool, its version family as pinned in package.json, and "
        "how it was used day-to-day.",
    )
    add_table(
        doc,
        ["Tool", "Version (repo)", "Purpose in this project"],
        [
            ("Node.js + npm", "18+ recommended", "Package management, running scripts"),
            ("Vite", "8.x", "Dev server with hot module replacement (HMR); production bundling to dist/"),
            ("React", "19.x", "Component-based UI, hooks for local state (useState, useMemo, useCallback)"),
            ("React DOM", "19.x", "Mounts the app into #root via createRoot"),
            ("React Router DOM", "7.x", "URL-based navigation, nested layouts, NavLink active states"),
            ("ESLint", "9.x flat config", "Catches unused variables, enforces React Hooks rules"),
            ("date-fns", "4.x", "Date formatting in SLT scheduling flows"),
            ("react-day-picker", "10.x", "Calendar UI inside Schedule Interview modal"),
            ("Google Fonts (Inter)", "CDN in index.html", "Primary typeface at weights 400–700"),
            ("Browser DevTools", "Chrome / Edge", "Layout inspection, accessibility checks, responsive preview"),
        ],
    )

    add_heading(doc, "3.1 IDE and project commands", 2)
    add_para(doc, "Source editing was done in a modern JavaScript IDE with JSX support and ESLint integration. Standard npm scripts:")
    add_code(doc, "npm install     # install dependencies\nnpm run dev     # start Vite at http://localhost:5173\nnpm run build   # production build → dist/\nnpm run preview # serve production build locally\nnpm run lint    # run ESLint across src/")
    add_para(
        doc,
        "Vite was chosen because layout work requires tight feedback loops: changing a CSS file "
        "or JSX file updates the browser in milliseconds without full page reloads for most edits.",
    )

    add_heading(doc, "3.2 What was intentionally not used", 2)
    add_bullet(doc, "No TypeScript — faster iteration for UI-only delivery; JSDoc comments document props")
    add_bullet(doc, "No CSS-in-JS (styled-components, etc.) — styles live in plain .css files per feature")
    add_bullet(doc, "No global state library (Redux, Zustand) — each page owns its state; layouts pass search via Router outlet context")
    add_bullet(doc, "No component library (MUI, Chakra) — custom components match the product design exactly")

    # 4 How UI works
    add_heading(doc, "4. How the User Interface Works", 1)

    add_heading(doc, "4.1 Application bootstrap", 2)
    add_para(doc, "When a user opens the app, the browser loads index.html, pulls Inter from Google Fonts, and executes src/main.jsx:")
    add_code(
        doc,
        "createRoot(#root).render(\n"
        "  <StrictMode>\n"
        "    <BrowserRouter>\n"
        "      <App />   ← all routes defined here\n"
        "    </BrowserRouter>\n"
        "  </StrictMode>\n"
        ")",
    )
    add_para(
        doc,
        "StrictMode enables React development checks. BrowserRouter synchronises the visible "
        "screen with the URL so bookmarks and back/forward navigation work as expected.",
    )

    add_heading(doc, "4.2 Routing and nested layouts", 2)
    add_para(
        doc,
        "App.jsx declares every route. Public routes render full-page login components. "
        "Authenticated-style portals use a parent Route with a Layout component and child routes "
        "that render inside <Outlet />.",
    )
    add_table(
        doc,
        ["URL prefix", "Layout component", "Example child routes"],
        [
            ("/", "LoginPage", "(none — full page)"),
            ("/staff-login", "StaffLoginPage", "(none)"),
            ("/parent", "ParentLayout", "dashboard, enrol-student, applications"),
            ("/admin", "AdminLayout", "dashboard, users, rules, form-builder, settings, …"),
            ("/slt", "SltLayout", "dashboard, review-queue, interviews, reports"),
            ("/manager", "ManagerLayout", "dashboard, applications, task-queue, assignments, reports"),
        ],
    )
    add_para(
        doc,
        "Each layout mounts once per portal visit. Only the outlet content swaps when the user "
        "clicks sidebar links—sidebar and top bar state (collapsed, search text) persist across pages.",
    )

    add_heading(doc, "4.3 Shared application shell (how every portal feels the same)", 2)
    add_para(doc, "All four portals use the same DOM structure and CSS foundation (parent-shell.css):")
    add_code(
        doc,
        "<div className=\"parent-app-shell\">\n"
        "  <AppSidebar brand={...} items={...} user={...} basePath=\"/admin\" />\n"
        "  <div className=\"parent-app-shell__column\">\n"
        "    <AppTopBar searchQuery={...} onSearchChange={...} />\n"
        "    <div className=\"parent-app-shell__scroll\">\n"
        "      <Outlet context={{ searchQuery }} />  ← page content\n"
        "    </div>\n"
        "  </div>\n"
        "</div>",
    )
    add_para(doc, "Responsibilities split as follows:")
    add_bullet(doc, "AppSidebar — brand, role pill, navigation links (React Router NavLink), user profile, sign-out")
    add_bullet(doc, "AppTopBar — global search input, MoE Verified badge, EN/Te Reo toggle, notification bell")
    add_bullet(doc, "Outlet — the active page component (dashboard, table, wizard, etc.)")
    add_para(
        doc,
        "Navigation items are not hard-coded inside AppSidebar. Each layout imports a config file "
        "(e.g. config/navigation/adminNav.js) and passes brand, roleLabel, and items as props. "
        "Adding a new admin section means updating the config array and adding one Route in App.jsx—"
        "the sidebar component itself stays unchanged.",
    )

    add_heading(doc, "4.4 Cross-page search (outlet context pattern)", 2)
    add_para(
        doc,
        "The top bar search field lives in the layout because it must appear on every page. "
        "Child pages that need to filter lists read the same value via React Router’s outlet context:",
    )
    add_code(
        doc,
        "// In layout (e.g. SltLayout / ManagerLayout):\n"
        "<Outlet context={{ searchQuery }} />\n\n"
        "// In page (e.g. SltReviewQueuePage):\n"
        "const { searchQuery } = useOutletContext()\n"
        "const filtered = queue.filter(item => matchesSearch(item, searchQuery))",
    )
    add_para(
        doc,
        "This avoids duplicating search UI on each page while keeping filter logic close to the "
        "data each page owns (queue items, table rows, etc.).",
    )

    add_heading(doc, "4.5 Page-level state and data flow", 2)
    add_para(doc, "Typical data flow within a screen:")
    add_numbered(doc, "Dummy or API data is imported from src/data/*.js (or will be fetched later)")
    add_numbered(doc, "Page component holds UI state with useState (selected row, open modal, active tab)")
    add_numbered(doc, "Derived lists use useMemo (filtered queue, selected item object)")
    add_numbered(doc, "Event handlers use useCallback when passed to children to avoid unnecessary re-renders")
    add_numbered(doc, "Presentational components receive props only—they do not fetch data themselves")
    add_para(
        doc,
        "Example: SltReviewQueuePage keeps the queue array in state, filters it with layout search, "
        "passes the selected application into SltApplicationReviewDetail, and opens "
        "ScheduleInterviewModal or RequestAdditionalInfoModal by setting an item ID in state.",
    )

    add_heading(doc, "4.6 Parent enrolment wizard flow", 2)
    add_para(
        doc,
        "ParentEnrolStudentPage orchestrates a seven-step wizard defined in config/parentEnrolment.js. "
        "The page owns one state object per step (studentDetails, caregiverDetails, …) and a stepIndex. "
        "Each step is a separate component implementing the same contract:",
    )
    add_code(
        doc,
        "<StudentDetailsStep\n"
        "  value={studentDetails}\n"
        "  onChange={(patch) => setStudentDetails(prev => ({ ...prev, ...patch }))}\n"
        "  yearLevelOptions={yearLevelOptions}\n"
        "/>",
    )
    add_para(
        doc,
        "Steps call onChange with a partial patch; the page merges into the full object. "
        "Progress bar width is (stepIndex + 1) / totalSteps. This pattern allows each step file "
        "to stay small (~150–250 lines) and makes it straightforward to add validation or "
        "save-draft API calls per step later.",
    )

    add_heading(doc, "4.7 Styling load order", 2)
    add_para(
        doc,
        "Global baseline: index.css (reset, body font, visually-hidden utility). Login pages use App.css. "
        "Each role layout imports parent-shell.css plus only the CSS bundles for that portal—for example "
        "AdminLayout imports admin-dashboard.css, admin-users.css, admin-rules.css, etc., so admin-specific "
        "rules are available on every admin page without re-importing per screen.",
    )

    # 5 Component architecture
    add_heading(doc, "5. Component Architecture and Reusability", 1)
    add_para(
        doc,
        "Components are grouped by responsibility, not by page. Pages should read as short "
        "composition layers: import widgets, map data to props, wire callbacks.",
    )

    add_heading(doc, "5.1 Folder structure", 2)
    add_table(
        doc,
        ["Folder", "Role", "Count (approx.)"],
        [
            ("src/components/layout/", "App shell shared by all portals", "2 components"),
            ("src/components/ui/", "Generic primitives (buttons, inputs, grids, toasts)", "8 components"),
            ("src/components/dashboard/", "Metrics, charts, tables, pills for operational UI", "25+ components"),
            ("src/components/enrolment/", "Parent wizard steps + EnrolSelect", "8+ step modules"),
            ("src/components/slt/", "Review queue, detail panel, modals", "10+ components"),
            ("src/components/icons/", "SVG icons + NavIcon name registry", "NavIcons.jsx, NavIcon.jsx"),
            ("src/config/navigation/", "Per-role sidebar config", "4 files"),
            ("src/config/", "Wizard steps, status labels, filter defaults", "3+ files"),
            ("src/data/", "Dummy JSON-like structures for demos", "10+ modules"),
            ("src/pages/", "Route-level screens only", "4 role folders + login pages"),
            ("src/styles/", "Feature-scoped CSS", "25+ stylesheets"),
        ],
    )

    add_heading(doc, "5.2 Reusability strategy", 2)
    add_para(doc, "Three deliberate patterns make components reusable across the system:")
    add_para(doc, "1. Configuration over duplication", bold=True)
    add_para(
        doc,
        "Navigation, enrolment step labels, and status pill labels live in src/config/. "
        "Components consume config at render time so the same EnrollmentStatusPill shows "
        "'SLT Review' on the manager table and the SLT queue without copying strings.",
    )
    add_para(doc, "2. Props-driven presentational components", bold=True)
    add_para(
        doc,
        "Dashboard widgets do not know which role is viewing them. StatSummaryCard only needs "
        "value, label, icon, and tone. DashboardPanel only needs title, optional action link, and children.",
    )
    add_para(doc, "3. Composition over inheritance", bold=True)
    add_para(
        doc,
        "Higher-level panels compose lower-level ones. RecentApplicationsPanel wraps DashboardPanel "
        "and maps rows to StatusPill. BarChartPanel and DonutBreakdownPanel wrap DashboardPanel for "
        "consistent headers and padding. No class inheritance is used anywhere in the codebase.",
    )

    add_heading(doc, "5.3 Layout components (used 4× — 100% of portals)", 2)
    add_table(
        doc,
        ["Component", "Key props", "Reused by"],
        [
            ("AppSidebar", "brand, roleLabel, items[], user, basePath, collapsed, onToggleCollapse, onSignOut", "Parent, Admin, SLT, Manager layouts"),
            ("AppTopBar", "searchPlaceholder, searchQuery, onSearchChange, notificationCount, demoTag", "All four layouts (demoTag on Manager only)"),
        ],
    )
    add_para(
        doc,
        "AppSidebar maps items[] to NavLink elements. Icons resolve through NavIcon, which looks "
        "up a string key (e.g. 'dashboard', 'shield') in a registry object—adding a nav icon means "
        "exporting one SVG component and adding one line to the byName map in NavIcon.jsx.",
    )

    add_heading(doc, "5.4 UI primitives (src/components/ui/)", 2)
    add_table(
        doc,
        ["Component", "Props / behaviour", "Used on"],
        [
            ("Button", "children, type, onClick, className", "Login pages, forms, modals"),
            ("InputField", "id, label, type, placeholder", "Login pages"),
            ("Card", "children, className", "Login form container"),
            ("FeatureItem", "title, description, icon", "Parent login marketing column"),
            ("DetailFieldGrid", "rows[{ id, label, value, valuePrefix? }]", "SLT application detail, review screens"),
            ("PanelEmptyState", "icon, title, description", "SLT queue when nothing selected / empty list"),
            ("AppToast", "message, tone, onDismiss", "SLT actions (approve, schedule, request info)"),
        ],
    )

    add_heading(doc, "5.5 Dashboard components — cross-role reuse", 2)
    add_para(
        doc,
        "The table below shows where the same component appears in more than one portal, "
        "which is the primary evidence of reuse.",
    )
    add_table(
        doc,
        ["Component", "Purpose", "Roles / pages using it"],
        [
            ("StatSummaryCard", "Numeric KPI with icon and colour tone", "Parent dashboard"),
            ("ApplicationListPanel + ApplicationRow", "List of applications with progress", "Parent dashboard, SLT home"),
            ("InfoActionCard", "Help / document call-to-action card", "Parent dashboard"),
            ("DashboardPanel", "Titled section with optional 'View all' link", "Admin users, charts, violations, integration"),
            ("AdminMetricCard, BarChartPanel, DonutBreakdownPanel", "Admin home analytics", "Admin dashboard only (composed from DashboardPanel)"),
            ("UserManagementTable, RoleSummaryCard", "User admin table", "Admin users"),
            ("ManagerKpiCard, ManagerApplicationQueueTable", "Manager workload", "Manager dashboard"),
            ("ManagerApplicationsTable", "Filterable applications grid", "Manager applications page"),
            ("EnrollmentStatusPill", "Status badge from statusKey + shared labels", "Manager tables, SLT queue, SLT detail, resolved card"),
            ("EnrollmentPriorityPill", "Priority badge", "Manager task queue"),
            ("SltWorkloadMemberCard", "SLT member workload summary", "Manager dashboard, assignments"),
            ("RuleCategoryCard, RuleConfigCard", "Rules engine UI", "Admin rules page"),
        ],
    )

    add_heading(doc, "5.6 EnrollmentStatusPill — example of config + CSS reuse", 2)
    add_para(
        doc,
        "Status display is centralised so every screen shows consistent labels and colours:",
    )
    add_code(
        doc,
        "// config/enrollmentApplicationUi.js\n"
        "enrollmentStatusLabels = { draft, submitted, sltReview, approved, ... }\n\n"
        "// EnrollmentStatusPill.jsx\n"
        "text = label ?? enrollmentStatusLabels[statusKey]\n"
        "className = `enrollment-status-pill--${statusKey}`\n\n"
        "// Used with only: <EnrollmentStatusPill statusKey={row.statusKey} />",
    )
    add_para(
        doc,
        "The same component is imported in manager tables, SLT list items, SLT detail header, "
        "and post-decision cards—six files, one source of truth for status styling (enrollment-pills.css).",
    )

    add_heading(doc, "5.7 SLT-specific composition", 2)
    add_para(
        doc,
        "SLT features use domain components that still depend on shared primitives:",
    )
    add_bullet(doc, "SltReviewQueueListItem — list row; uses EnrollmentStatusPill")
    add_bullet(doc, "SltApplicationReviewDetail — tabbed detail; uses DetailFieldGrid, EnrollmentStatusPill, document/timeline/notes sub-panels")
    add_bullet(doc, "ScheduleInterviewModal — react-day-picker + date-fns; isolated CSS file")
    add_bullet(doc, "RequestAdditionalInfoModal — email template from data/requestAdditionalInfoEmailTemplate.js")
    add_para(
        doc,
        "SltApplicationReviewDetail exposes callback props (onApprove, onDecline, onRequestInfo, …) "
        "so SltReviewQueuePage owns queue mutations and toast feedback—the detail panel stays a pure UI surface.",
    )

    add_heading(doc, "5.8 Enrolment step components", 2)
    add_table(
        doc,
        ["Step component", "Config id", "State slice owned by page"],
        [
            ("StudentDetailsStep", "student", "studentDetails"),
            ("CaregiverInformationStep", "caregiver", "caregiverDetails"),
            ("AddressAndZoningStep", "address", "addressDetails"),
            ("PreviousSchoolStep", "previous", "previousSchoolDetails"),
            ("MedicalInformationStep", "medical", "medicalDetails"),
            ("DocumentsStep", "documents", "documentsDetails"),
            ("ReviewSubmitStep", "review", "reviewDetails"),
        ],
    )
    add_para(
        doc,
        "EnrolSelect is shared by steps that need dropdowns. Default field values live in "
        "*Defaults.js files next to each step so the page’s useState initialisers stay readable.",
    )

    # 6 Technology reference
    add_heading(doc, "6. Technology Stack Reference", 1)
    add_table(
        doc,
        ["Layer", "Choice", "Rationale"],
        [
            ("UI library", "React 19", "Mature component model; hooks for forms and wizards"),
            ("Bundler / dev server", "Vite 8", "Fast HMR for visual work"),
            ("Routing", "React Router 7", "Nested routes match nested layouts"),
            ("Styling", "Plain CSS per feature", "Easy overrides; no runtime style injection"),
            ("Icons", "Inline SVG components", "Crisp at any DPI; colour via currentColor"),
            ("Dates", "date-fns + react-day-picker", "Interview scheduling modal only"),
        ],
    )

    # 7 Design system
    add_heading(doc, "7. Design System", 1)
    add_heading(doc, "7.1 Typography and colour tokens", 2)
    add_bullet(doc, "Font: Inter (400–700) loaded in index.html")
    add_bullet(doc, "Body: #1f2937 on #f3f5f7 page background (parent-shell.css --page-bg)")
    add_bullet(doc, "Primary chrome: #163a5c sidebar (--sidebar-teal), active link #244f78")
    add_bullet(doc, "Cards: white background, 12–16px radius, borders #e5e7eb / #e2e8f0")
    add_heading(doc, "7.2 CSS naming convention", 2)
    add_para(
        doc,
        "Classes use BEM-style prefixes to avoid collisions: app-sidebar__, parent-enrol__, "
        "slt-rq-, enrollment-status-pill--{statusKey}. Feature CSS files are imported only "
        "where needed, keeping the global cascade predictable.",
    )
    add_heading(doc, "7.3 Bilingual navigation", 2)
    add_para(
        doc,
        "Each nav item in config/navigation/*.js supports label (English) and subLabel (te reo Māori). "
        "AppSidebar renders both lines when the sidebar is expanded; collapsed mode shows icons only with title tooltips.",
    )

    # 8 Screenshots
    add_heading(doc, "8. Screen Reference (with figures)", 1)

    add_heading(doc, "8.1 Authentication", 2)
    add_image(doc, "01-login.png", "Figure 1 — Parent login: FeatureItem cards + InputField form")
    add_image(doc, "02-staff-login.png", "Figure 2 — Staff login with portal demo links")

    add_heading(doc, "8.2 Parent portal", 2)
    add_para(
        doc,
        "Dashboard composes StatSummaryCard, ApplicationListPanel, and InfoActionCard. "
        "Enrolment route renders the wizard with step indicator and shared shell.",
    )
    add_image(doc, "03-parent-dashboard.png", "Figure 3 — Parent dashboard composition")
    add_image(doc, "04-parent-enrolment.png", "Figure 4 — Seven-step enrolment wizard")

    add_heading(doc, "8.3 Administrator portal", 2)
    add_image(doc, "05-admin-dashboard.png", "Figure 5 — Admin dashboard (metric cards + chart panels)")
    add_image(doc, "06-admin-users.png", "Figure 6 — Users page (RoleSummaryCard + DashboardPanel + table)")

    add_heading(doc, "8.4 SLT portal", 2)
    add_image(doc, "07-slt-dashboard.png", "Figure 7 — SLT dashboard")
    add_image(doc, "08-slt-review-queue.png", "Figure 8 — Master-detail review queue")

    add_heading(doc, "8.5 Manager portal", 2)
    add_image(doc, "09-manager-dashboard.png", "Figure 9 — Manager KPI and queue widgets")
    add_image(doc, "10-manager-applications.png", "Figure 10 — Applications table with EnrollmentStatusPill")

    # 9 Integration
    add_heading(doc, "9. Build, Quality, and Back-End Integration", 1)
    add_heading(doc, "9.1 Build pipeline", 2)
    add_code(doc, "npm run build  →  Vite bundles JS/CSS/assets  →  output in dist/\nnpm run preview  →  local static server for UAT")
    add_heading(doc, "9.2 Integration points (ready without refactor)", 2)
    add_bullet(doc, "Replace src/data/*.js imports with fetch/axios hooks returning the same object shapes")
    add_bullet(doc, "Wrap layouts with an auth context; redirect unauthenticated users from /parent, /admin, etc.")
    add_bullet(doc, "Persist wizard state: POST partial step payloads from each onChange or on Next click")
    add_bullet(doc, "Wire AppTopBar language toggle to i18n provider when translations are available")

    add_heading(doc, "9.3 Accessibility measures in place", 2)
    add_bullet(doc, "aria-label on sidebar, search, notifications, collapse toggle")
    add_bullet(doc, "aria-pressed on language pills")
    add_bullet(doc, "Semantic headings inside dashboard panels")
    add_bullet(doc, "visually-hidden utility class in index.css for screen-reader-only text")

    # Appendix
    add_heading(doc, "Appendix A — Revision history", 1)
    add_table(
        doc,
        ["Version", "Date", "Changes"],
        [
            ("1.0", DOC_DATE, "Initial documentation with screenshots"),
            ("1.1", DOC_DATE, "Expanded: tools, UI runtime behaviour, component reuse tables"),
        ],
    )

    add_heading(doc, "Appendix B — Key source files quick reference", 1)
    add_table(
        doc,
        ["File", "Responsibility"],
        [
            ("src/main.jsx", "App bootstrap + BrowserRouter"),
            ("src/App.jsx", "All route definitions"),
            ("src/pages/*/ *Layout.jsx", "Shell + outlet per role"),
            ("src/config/navigation/*.js", "Sidebar items per role"),
            ("src/components/layout/AppSidebar.jsx", "Shared sidebar"),
            ("src/components/layout/AppTopBar.jsx", "Shared header bar"),
            ("src/styles/parent-shell.css", "Shell layout + sidebar/topbar styles"),
            ("src/config/enrollmentApplicationUi.js", "Shared status/priority labels"),
        ],
    )

    doc.save(OUTPUT)
    print(f"Written: {OUTPUT}")


if __name__ == "__main__":
    build_document()
